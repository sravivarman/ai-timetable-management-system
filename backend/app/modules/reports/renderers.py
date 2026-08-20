"""In-memory XLSX, CSV, DOCX and PDF renderers for canonical reports."""

from __future__ import annotations

import csv
import re
from datetime import datetime, timezone
from io import BytesIO, StringIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import LongTable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, TableStyle

from app.modules.reports.service import CanonicalReport


MIME_TYPES = {
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def render_report(result: CanonicalReport, export_format: str) -> bytes:
    if result.definition.layout_type == "SLOT_TIMETABLE" and export_format != "csv":
        renderer = {"xlsx": render_slot_xlsx, "docx": render_slot_docx, "pdf": render_slot_pdf}.get(export_format)
        if renderer is not None:
            return renderer(result)
    renderer = {"xlsx": render_xlsx, "csv": render_csv, "docx": render_docx, "pdf": render_pdf}.get(export_format)
    if renderer is None:
        raise ValueError("Unsupported report format")
    return renderer(result)


def report_filename(result: CanonicalReport, export_format: str) -> str:
    context = next((item.split(":", 1)[1].strip() for item in result.filter_summary if item.startswith("Academic Term:")), None)
    value = "_".join(item for item in (result.definition.title, context) if item)
    return f"{re.sub(r'[^A-Za-z0-9_-]+', '_', value).strip('_')}.{export_format}"


def render_csv(result: CanonicalReport) -> bytes:
    stream = StringIO(newline="")
    writer = csv.writer(stream)
    writer.writerow([column.label for column in result.columns])
    for row in result.rows:
        writer.writerow(["" if row.get(column.key) is None else row.get(column.key) for column in result.columns])
    return ("\ufeff" + stream.getvalue()).encode("utf-8")


def render_xlsx(result: CanonicalReport) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = re.sub(r"[\\/*?:\[\]]", "", result.definition.title)[:31] or "Report"
    sheet.cell(1, 1, result.definition.title).font = Font(size=16, bold=True, color="17365D")
    sheet.cell(2, 1, "Filters: " + ("; ".join(result.filter_summary) or "None"))
    sheet.cell(3, 1, f"Generated: {datetime.now(timezone.utc).isoformat()}")
    header_row = 5
    for index, column in enumerate(result.columns, 1):
        cell = sheet.cell(header_row, index, column.label)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.column_dimensions[get_column_letter(index)].width = max(12, min(column.default_width, 42))
    rows = result.rows or ({column.key: "No records found" if index == 0 else None for index, column in enumerate(result.columns)},)
    for row_index, row in enumerate(rows, header_row + 1):
        for column_index, column in enumerate(result.columns, 1):
            value = row.get(column.key)
            cell = sheet.cell(row_index, column_index, value)
            cell.alignment = Alignment(horizontal=column.alignment, vertical="top", wrap_text=True)
    sheet.freeze_panes = f"A{header_row + 1}"
    sheet.auto_filter.ref = f"A{header_row}:{get_column_letter(len(result.columns))}{header_row + len(rows)}"
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def render_slot_xlsx(result: CanonicalReport) -> bytes:
    """Render an actual-date matrix plus the canonical long-form detail sheet."""
    workbook = Workbook()
    matrix = workbook.active
    matrix.title = "Timetable Matrix"
    matrix.append([result.definition.title])
    matrix.append(["Filters", "; ".join(result.filter_summary) or "None"])
    matrix.append([])
    matrix.append(["Date", "Day", "P1", "P2", "P3", "P4", "P5", "P6", "P7"])
    for cell in matrix[4]:
        cell.font = Font(bold=True, color="FFFFFF"); cell.fill = PatternFill("solid", fgColor="1F4E78"); cell.alignment = Alignment(horizontal="center", wrap_text=True)
    for date_value, day, periods in _slot_matrix_rows(result):
        matrix.append([date_value, day, *periods])
    matrix.freeze_panes = "C5"
    matrix.column_dimensions["A"].width = 14; matrix.column_dimensions["B"].width = 12
    for index in range(3, 10): matrix.column_dimensions[get_column_letter(index)].width = 28
    for row in matrix.iter_rows(min_row=5):
        for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)

    detail = workbook.create_sheet("Long-form Detail")
    for index, column in enumerate(result.columns, 1):
        cell = detail.cell(1, index, column.label); cell.font = Font(bold=True, color="FFFFFF"); cell.fill = PatternFill("solid", fgColor="1F4E78")
        detail.column_dimensions[get_column_letter(index)].width = max(12, min(column.default_width, 42))
    for row_index, row in enumerate(result.rows, 2):
        for column_index, column in enumerate(result.columns, 1): detail.cell(row_index, column_index, row.get(column.key)).alignment = Alignment(vertical="top", wrap_text=True)
    detail.freeze_panes = "A2"
    if result.columns: detail.auto_filter.ref = f"A1:{get_column_letter(len(result.columns))}{max(1, len(result.rows)+1)}"
    stream = BytesIO(); workbook.save(stream); return stream.getvalue()


def render_slot_docx(result: CanonicalReport) -> bytes:
    document = Document(); section = document.sections[0]; section.orientation = WD_ORIENT.LANDSCAPE; section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("AI Timetable Management System", 0); document.add_heading(result.definition.title, level=1)
    document.add_paragraph("Filters: " + ("; ".join(result.filter_summary) or "None"))
    table = document.add_table(rows=1, cols=9); table.style = "Table Grid"
    for index, label in enumerate(("Date", "Day", "P1", "P2", "P3", "P4", "P5", "P6", "P7")):
        table.rows[0].cells[index].text = label
        for run in table.rows[0].cells[index].paragraphs[0].runs: run.bold = True; run.font.size = Pt(8)
    table.rows[0]._tr.get_or_add_trPr().append(_repeat_table_header())
    for date_value, day, periods in _slot_matrix_rows(result):
        cells = table.add_row().cells
        for index, value in enumerate((date_value, day, *periods)):
            cells[index].text = value
            for run in cells[index].paragraphs[0].runs: run.font.size = Pt(7)
    stream = BytesIO(); document.save(stream); return stream.getvalue()


def render_slot_pdf(result: CanonicalReport) -> bytes:
    pagesize = landscape(A4); stream = BytesIO(); document = SimpleDocTemplate(stream, pagesize=pagesize, leftMargin=8*mm, rightMargin=8*mm, topMargin=10*mm, bottomMargin=12*mm)
    styles = getSampleStyleSheet(); cell = ParagraphStyle("SlotCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=6.2, leading=7.2); header = ParagraphStyle("SlotHeader", parent=cell, textColor=colors.white, alignment=TA_CENTER, fontName="Helvetica-Bold")
    story: list[Any] = [Paragraph("AI Timetable Management System", styles["Title"]), Paragraph(result.definition.title, styles["Heading1"]), Paragraph("Filters: " + escape("; ".join(result.filter_summary) or "None"), styles["BodyText"]), Spacer(1, 4*mm)]
    labels = ("Date", "Day", "P1", "P2", "P3", "P4", "P5", "P6", "P7")
    data = [[Paragraph(label, header) for label in labels]]
    for date_value, day, periods in _slot_matrix_rows(result): data.append([Paragraph(escape(value), cell) for value in (date_value, day, *periods)])
    widths = [24*mm, 18*mm] + [32*mm]*7; table = LongTable(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4E78")),("GRID",(0,0),(-1,-1),0.35,colors.HexColor("#9CA3AF")),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),2),("RIGHTPADDING",(0,0),(-1,-1),2),("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2)]))
    story.append(table); document.build(story, onFirstPage=_page_number, onLaterPages=_page_number); return stream.getvalue()


def _slot_matrix_rows(result: CanonicalReport) -> list[tuple[str, str, list[str]]]:
    grouped: dict[tuple[str, str], dict[int, list[str]]] = {}
    for row in result.rows:
        key = (str(row.get("date") or "Unscheduled"), str(row.get("day") or "")); period = int(row.get("period") or 0)
        if not 1 <= period <= 7: continue
        duration = max(1, int(row.get("duration") or 1)); span = f"P{period}" if duration == 1 else f"P{period}–P{period+duration-1}"
        lines = [f"{row.get('course_code') or 'Course'} · {row.get('section_code') or ''}", str(row.get("student_group") or ""), str(row.get("faculty_code") or ""), str(row.get("venue_code") or ""), span]
        label = "\n".join(value for value in lines if value)
        values = grouped.setdefault(key, {}).setdefault(period, [])
        if label not in values: values.append(label)
    if not grouped: return [("No records", "", [""]*7)]
    return [(date_value, day, ["\n\n".join(grouped[(date_value, day)].get(period, [])) for period in range(1,8)]) for date_value, day in sorted(grouped)]


def render_docx(result: CanonicalReport) -> bytes:
    document = Document()
    if len(result.columns) > 6 or sum(column.default_width for column in result.columns) > 115:
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("AI Timetable Management System", 0)
    document.add_heading(result.definition.title, level=1)
    document.add_paragraph("Filters: " + ("; ".join(result.filter_summary) or "None"))
    document.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
    table = document.add_table(rows=1, cols=len(result.columns))
    table.style = "Table Grid"
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(_repeat_table_header())
    for index, column in enumerate(result.columns):
        header.cells[index].text = column.label
        for run in header.cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        header.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = result.rows or ({result.columns[0].key: "No records found"},)
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(result.columns):
            cells[index].text = display_value(row.get(column.key))
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(8)
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer.add_run("Generated by AI Timetable Management System")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _repeat_table_header():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    return element


def render_pdf(result: CanonicalReport) -> bytes:
    wide = len(result.columns) > 6 or sum(column.default_width for column in result.columns) > 115
    pagesize = landscape(A4) if wide else A4
    stream = BytesIO()
    document = SimpleDocTemplate(stream, pagesize=pagesize, leftMargin=10 * mm, rightMargin=10 * mm, topMargin=12 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("ReportCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=6.5 if wide else 8, leading=8 if wide else 10)
    header_style = ParagraphStyle("ReportHeader", parent=cell_style, textColor=colors.white, alignment=TA_CENTER, fontName="Helvetica-Bold")
    story: list[Any] = [Paragraph("AI Timetable Management System", styles["Title"]), Paragraph(result.definition.title, styles["Heading1"])]
    story.append(Paragraph("Filters: " + escape("; ".join(result.filter_summary) or "None"), styles["BodyText"]))
    story.append(Paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}", styles["BodyText"]))
    story.append(Spacer(1, 5 * mm))
    data = [[Paragraph(escape(column.label), header_style) for column in result.columns]]
    if result.rows:
        for row in result.rows:
            data.append([Paragraph(escape(display_value(row.get(column.key))), cell_style) for column in result.columns])
    else:
        data.append([Paragraph("No records found", cell_style)] + [Paragraph("", cell_style) for _ in result.columns[1:]])
    available = pagesize[0] - 20 * mm
    weights = [max(10, min(column.default_width, 34)) for column in result.columns]
    total_weight = sum(weights)
    widths = [available * weight / total_weight for weight in weights]
    table = LongTable(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3), ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F7FA")]),
    ]))
    story.append(table)
    document.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    return stream.getvalue()


def _page_number(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(document.pagesize[0] - 10 * mm, 7 * mm, f"Page {document.page}")
    canvas.restoreState()


def display_value(value: Any) -> str:
    return "—" if value is None or value == "" else str(value)


def escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
