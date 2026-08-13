"""Administrative report datasets, validation, exports, and authorization."""

import csv
import io
import os
import re
import unittest
from datetime import date

os.environ.setdefault("DATABASE_URL", "postgresql+psycopg://postgres:postgres@localhost:5432/timetable_db")
os.environ.setdefault("SECRET_KEY", "test-secret-that-is-at-least-thirty-two-bytes")

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import select

from app.main import app
from app.modules.academic_terms.models import AcademicTerm
from app.modules.authentication.models import Permission, Role
from app.modules.course_offerings.models import CourseOffering, CourseOfferingAllowedLaboratory
from app.modules.courses.models import Course, CourseEligibleLaboratory
from app.modules.facilities.models import Classroom, Laboratory
from app.modules.facilities_constraints.models import SectionClassroomAssignment
from app.modules.faculty.models import Faculty
from app.modules.faculty_allocations.models import LaboratoryFacultyAllocation, TheoryFacultyAllocation
from app.modules.faculty_allocations.workload import configured_faculty_workloads
from app.modules.programs.models import Program
from app.modules.sections.models import Section
from tests.facilities_test_support import create_facilities_test_context


class AdministrativeReportTests(unittest.TestCase):
    def setUp(self):
        self.context = create_facilities_test_context()
        db = self.context.session_factory()
        try:
            permission = Permission(resource="reports", action="read", description="Read reports")
            db.add(permission); db.flush()
            for role in db.scalars(select(Role).where(Role.name.in_(["Administrator", "Timetable Coordinator", "HOD"]))):
                role.permissions.append(permission)
            term = AcademicTerm(academic_year="2026-27", term_name="I-I", year_number=1, semester_number=1, start_date=date(2026, 7, 1), end_date=date(2026, 11, 30), is_active=True, is_current=True)
            program = Program(department_id=self.context.active_department.id, program_code="BTECH-TST", program_name="B.Tech Test")
            db.add_all([term, program]); db.flush()
            first = Section(program_id=program.id, academic_term_id=term.id, section_name="A", section_code="TST-A", student_strength=72)
            second = Section(program_id=program.id, academic_term_id=term.id, section_name="B", section_code="TST-B", student_strength=60)
            theory = Course(course_code="TST101", course_name="Theory Course", offering_department_id=self.context.active_department.id, course_type="THEORY", weekly_periods=4, counts_toward_workload=True)
            activity = Course(course_code="TST102", course_name="Practical Course", offering_department_id=self.context.active_department.id, course_type="PRACTICAL", weekly_periods=2, grouping_mode="GROUPED", default_group_count=2, session_duration=2, sessions_per_week=1, venue_requirement="LABORATORY_ONLY", counts_toward_workload=True)
            missing = Course(course_code="TST103", course_name="Unallocated Course", offering_department_id=self.context.active_department.id, course_type="THEORY", weekly_periods=3, counts_toward_workload=True)
            room = Classroom(room_number="R101", room_name="Test Classroom", owning_department_id=self.context.active_department.id)
            laboratory = Laboratory(laboratory_code="TST-LAB", laboratory_name="Test Laboratory", room_number="L101", owning_department_id=self.context.active_department.id)
            main = Faculty(faculty_code="VCE101", full_name="Main Faculty", department_id=self.context.active_department.id, designation="Professor", institutional_email="main.report@vce.ac.in", minimum_weekly_workload=4, maximum_weekly_workload=20)
            supporting = Faculty(faculty_code="VCE102", full_name="Supporting Faculty", department_id=self.context.active_department.id, designation="Assistant Professor", institutional_email="support.report@vce.ac.in", minimum_weekly_workload=8, maximum_weekly_workload=20)
            zero = Faculty(faculty_code="VCE103", full_name="Zero Faculty", department_id=self.context.active_department.id, designation="Associate Professor", institutional_email="zero.report@vce.ac.in", minimum_weekly_workload=6, maximum_weekly_workload=18)
            db.add_all([first, second, theory, activity, missing, room, laboratory, main, supporting, zero]); db.flush()
            offerings = [
                CourseOffering(course_id=theory.id, section_id=first.id, academic_term_id=term.id),
                CourseOffering(course_id=theory.id, section_id=second.id, academic_term_id=term.id),
                CourseOffering(course_id=activity.id, section_id=first.id, academic_term_id=term.id, laboratory_selection_mode="RESTRICTED"),
                CourseOffering(course_id=missing.id, section_id=first.id, academic_term_id=term.id),
            ]
            db.add_all(offerings); db.flush()
            db.add_all([
                CourseEligibleLaboratory(course_id=activity.id, laboratory_id=laboratory.id),
                CourseOfferingAllowedLaboratory(course_offering_id=offerings[2].id, laboratory_id=laboratory.id),
                SectionClassroomAssignment(section_id=first.id, classroom_id=room.id, academic_term_id=term.id, is_primary=True),
                TheoryFacultyAllocation(course_offering_id=offerings[0].id, faculty_id=main.id),
                TheoryFacultyAllocation(course_offering_id=offerings[1].id, faculty_id=main.id),
                LaboratoryFacultyAllocation(course_offering_id=offerings[2].id, faculty_id=main.id, role_type="MAIN"),
                LaboratoryFacultyAllocation(course_offering_id=offerings[2].id, faculty_id=supporting.id, role_type="SUPPORTING", required_with_main_faculty_id=main.id),
            ])
            db.commit()
            self.term_id, self.department_id, self.section_id = term.id, self.context.active_department.id, first.id
            self.main_id, self.activity_offering_id = main.id, offerings[2].id
        finally:
            db.close()
        self.client = TestClient(app)

    def tearDown(self):
        self.client.close(); self.context.close()

    def request(self, report_key, columns, **filters):
        return {"report_key": report_key, "filters": filters, "selected_columns": columns, "sort_fields": [], "page": 1, "page_size": 100}

    def test_metadata_validation_and_authorization(self):
        response = self.client.get("/api/v1/reports/definitions", headers=self.context.headers["administrator"])
        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual([item["title"] for item in response.json()], ["Faculty Master", "Course Offerings", "Theory Faculty Allocations", "Activity Faculty Allocations", "Section-wise Course & Faculty Allocation", "Faculty Workload"])
        self.assertEqual(self.client.get("/api/v1/reports/definitions", headers=self.context.headers["coordinator"]).status_code, 200)
        self.assertEqual(self.client.get("/api/v1/reports/definitions", headers=self.context.headers["hod"]).status_code, 200)
        self.assertEqual(self.client.get("/api/v1/reports/definitions", headers=self.context.headers["unauthorized"]).status_code, 403)
        invalid = self.client.post("/api/v1/reports/preview", json=self.request("faculty_master", ["faculty_name", "not_a_column"]), headers=self.context.headers["administrator"])
        self.assertEqual(invalid.status_code, 422)
        invalid_filter = self.client.post("/api/v1/reports/preview", json=self.request("faculty_master", ["faculty_name"], academic_term_id=str(self.term_id)), headers=self.context.headers["administrator"])
        self.assertEqual(invalid_filter.status_code, 422)
        invalid_sort = self.request("faculty_master", ["faculty_name"]); invalid_sort["sort_fields"] = [{"key": "course_code", "direction": "asc"}]
        self.assertEqual(self.client.post("/api/v1/reports/preview", json=invalid_sort, headers=self.context.headers["administrator"]).status_code, 422)

    def test_faculty_master_filter_and_selected_column_order(self):
        payload = self.request("faculty_master", ["faculty_name", "faculty_code", "designation", "institutional_email"], department_id=str(self.department_id))
        response = self.client.post("/api/v1/reports/preview", json=payload, headers=self.context.headers["administrator"])
        self.assertEqual(response.status_code, 200, response.text)
        body = response.json()
        self.assertEqual([item["key"] for item in body["columns"]], payload["selected_columns"])
        self.assertEqual(body["total"], 3)
        self.assertEqual(list(body["rows"][0]), payload["selected_columns"])
        visible = str(body["rows"]) + str(body["filter_summary"])
        self.assertNotRegex(visible, r"[0-9a-f]{8}-[0-9a-f]{4}-[1-5][0-9a-f]{3}-[89ab][0-9a-f]{3}-[0-9a-f]{12}")

    def test_offering_theory_and_activity_semantics(self):
        offerings = self.client.post("/api/v1/reports/preview", json=self.request("course_offerings", ["course_code", "course_type", "laboratory_selection_mode", "allowed_laboratories"], academic_term_id=str(self.term_id)), headers=self.context.headers["administrator"]).json()
        practical = next(item for item in offerings["rows"] if item["course_code"] == "TST102")
        self.assertEqual(practical["course_type"], "PRACTICAL")
        self.assertEqual(practical["laboratory_selection_mode"], "Restricted")
        self.assertIn("TST-LAB", practical["allowed_laboratories"])
        theory = self.client.post("/api/v1/reports/preview", json=self.request("theory_faculty_allocations", ["section_code", "course_code", "faculty_code"], academic_term_id=str(self.term_id)), headers=self.context.headers["administrator"]).json()
        self.assertEqual(theory["total"], 2)
        self.assertEqual({item["faculty_code"] for item in theory["rows"]}, {"VCE101"})
        activity = self.client.post("/api/v1/reports/preview", json=self.request("activity_faculty_allocations", ["course_type", "faculty_code", "faculty_role", "required_main_faculty"], academic_term_id=str(self.term_id)), headers=self.context.headers["administrator"]).json()
        self.assertEqual(activity["total"], 2)
        self.assertEqual({item["faculty_role"] for item in activity["rows"]}, {"MAIN", "SUPPORTING"})
        supporting = next(item for item in activity["rows"] if item["faculty_role"] == "SUPPORTING")
        self.assertIn("VCE101", supporting["required_main_faculty"])

    def test_section_report_keeps_missing_and_activity_roles(self):
        payload = self.request("section_course_faculty", ["section_code", "course_code", "faculty_code", "faculty_role", "faculty_allocation_status", "primary_classroom"], section_id=str(self.section_id))
        body = self.client.post("/api/v1/reports/preview", json=payload, headers=self.context.headers["administrator"]).json()
        missing = next(item for item in body["rows"] if item["course_code"] == "TST103")
        self.assertEqual(missing["faculty_allocation_status"], "MISSING")
        activity_rows = [item for item in body["rows"] if item["course_code"] == "TST102"]
        self.assertEqual({item["faculty_role"] for item in activity_rows}, {"MAIN", "SUPPORTING"})
        self.assertTrue(all(item["faculty_allocation_status"] == "COMPLETE" for item in activity_rows))
        self.assertIn("R101", missing["primary_classroom"])

    def test_workload_includes_zero_and_matches_authoritative_service(self):
        payload = self.request("faculty_workload", ["faculty_code", "theory_workload", "activity_workload", "total_workload", "workload_status"], academic_term_id=str(self.term_id), department_id=str(self.department_id))
        body = self.client.post("/api/v1/reports/preview", json=payload, headers=self.context.headers["administrator"]).json()
        by_code = {row["faculty_code"]: row for row in body["rows"]}
        self.assertEqual(by_code["VCE103"]["total_workload"], 0)
        self.assertEqual(by_code["VCE103"]["workload_status"], "UNDERLOAD")
        db = self.context.session_factory()
        try: expected = configured_faculty_workloads(db, academic_term_id=self.term_id)
        finally: db.close()
        self.assertEqual(by_code["VCE101"]["total_workload"], expected[self.main_id])
        self.assertEqual(by_code["VCE101"]["total_workload"], by_code["VCE101"]["theory_workload"] + by_code["VCE101"]["activity_workload"])

    def test_all_export_formats_are_structurally_valid_and_consistent(self):
        columns = ["faculty_name", "faculty_code", "designation"]
        payload = self.request("faculty_master", columns, department_id=str(self.department_id))
        preview = self.client.post("/api/v1/reports/preview", json=payload, headers=self.context.headers["administrator"]).json()
        csv_response = self.client.post("/api/v1/reports/export?format=csv", json=payload, headers=self.context.headers["administrator"])
        rows = list(csv.reader(io.StringIO(csv_response.content.decode("utf-8-sig"))))
        self.assertEqual(rows[0], ["Faculty Name", "Faculty ID", "Designation"])
        self.assertEqual(len(rows) - 1, preview["total"])
        xlsx = self.client.post("/api/v1/reports/export?format=xlsx", json=payload, headers=self.context.headers["administrator"])
        sheet = load_workbook(io.BytesIO(xlsx.content)).active
        self.assertEqual([sheet.cell(5, index).value for index in range(1, 4)], rows[0])
        self.assertEqual(sheet.freeze_panes, "A6")
        self.assertEqual(sheet.max_row - 5, preview["total"])
        docx = self.client.post("/api/v1/reports/export?format=docx", json=payload, headers=self.context.headers["administrator"])
        document = Document(io.BytesIO(docx.content))
        self.assertIn("Faculty Master", " ".join(paragraph.text for paragraph in document.paragraphs))
        self.assertEqual([cell.text for cell in document.tables[0].rows[0].cells], rows[0])
        self.assertEqual(len(document.tables[0].rows) - 1, preview["total"])
        pdf = self.client.post("/api/v1/reports/export?format=pdf", json=payload, headers=self.context.headers["administrator"])
        self.assertTrue(pdf.content.startswith(b"%PDF")); self.assertGreater(len(pdf.content), 1000)
        self.assertIn("attachment", csv_response.headers["content-disposition"])


if __name__ == "__main__":
    unittest.main()
